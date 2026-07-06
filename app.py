import re
import json
from io import BytesIO
from pathlib import Path
from copy import deepcopy
from collections import Counter, defaultdict
from bs4 import BeautifulSoup, NavigableString, Tag, Comment
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn



# ============================================================
# GLOBAL CONFIG
# ============================================================

DEFAULT_FONT = "Arial"
DEFAULT_FONT_SIZE_PT = 10.5 

PAGE_MARGIN_TOP = 0.7
PAGE_MARGIN_BOTTOM = 0.7
PAGE_MARGIN_LEFT = 0.6
PAGE_MARGIN_RIGHT = 0.6

OUTPUT_TABLE_BORDER_COLOR = "000000"
OUTPUT_TABLE_BORDER_SIZE = "6"

HEADER_FILL_COLOR = "D9D9D9"

BASE_DIR = Path(__file__).resolve().parent
INPUT_HTML_FILE = BASE_DIR / "input.html"

#camel- to- snake case for json

# def camel_to_snake(name):
#     return re.sub(r'(?<!^)(?=[A-Z])', '_', name).lower()

# def convert_keys_to_snake_case(data):
#     if isinstance(data, dict):
#         return {
#             camel_to_snake(k): convert_keys_to_snake_case(v)
#             for k, v in data.items()
#         }
#     elif isinstance(data, list):
#         return [convert_keys_to_snake_case(i) for i in data]
#     else:
#         return data
    
# def convert_html_placeholders_to_snake_case(html: str) -> str:
#     def replace_match(match):
#         var = match.group(1).strip()
#         return "{{" + camel_to_snake(var) + "}}"

#     return re.sub(r"\{\{\s*(.*?)\s*\}\}", replace_match, html)
# #json cleaning
# def clean_json_keys(data):
#     if isinstance(data, dict):
#         new_dict = {}
#         for k, v in data.items():
#             clean_key = k.replace("{", "").replace("}", "")
#             new_dict[clean_key] = clean_json_keys(v)
#         return new_dict
#     elif isinstance(data, list):
#         return [clean_json_keys(i) for i in data]
#     else:
#         return data
# ============================================================
# FUNCTION 1:
# JSON INPUT + HTML TEMPLATE -> FILLED HTML STRING
# ============================================================

def json_to_filled_html_string(html_template: str, json_data: dict) -> str:
    """
    Takes:
        - html_template: HTML string containing placeholders like {{ variable_name }}
        - json_data: input JSON dictionary containing:
            json_data["llm_json"]
            json_data["mis_json"]

    Returns:
        - final filled HTML string

    Notes:
        - Does not save any HTML file.
        - Replaces scalar placeholders using LLM + MIS data.
        - Fills configured dynamic tables.
        - Any unreplaced {{ placeholder }} is replaced with 'rm_input'.
    """

    # ------------------------------------------------------------
    # Extract LLM and MIS sections from input JSON
    # ------------------------------------------------------------
    
    # html_template = convert_html_placeholders_to_snake_case(html_template)
    
    # json_data = clean_json_keys(json_data)

    # json_data = convert_keys_to_snake_case(json_data)

    #new change
    #if "mis_json" in json_data:
    #    json_data["mis_json"]["validity_of_limit"] = json_data["mis_json"].get("validityOfLimit")
    #if "mis_json" in json_data:
    #    json_data["mis_json"]["proposed_annual_credit_limit"] = json_data["mis_json"].get("creditLimit")

    datal = json_data["llm_json"]
    datax = json_data["mis_json"]

    def nest_dict():
        return defaultdict(nest_dict)
    
    datw = nest_dict()
    datw["web_results"]["result"] = datal["web_result"]
    dataw = dict(datw)
    
    data = dataw | datal
    # ------------------------------------------------------------
    # Convert MIS data into a mutable defaultdict
    # ------------------------------------------------------------

    mis_json = defaultdict(list)

    if "data" in datax:
        for key, value in datax["data"].items():
            mis_json[key.lower()] = value
    else:
        for key, value in datax.items():
            mis_json[key.lower()] = value
    
    # ------------------------------------------------------------
    # Prepare sanction table data from objDanMisDetails
    # ------------------------------------------------------------

    sanction_rows = []

    if "data" in datax:
        dt = datax["data"]
    else:
        dt = datax
        
    for item in dt.get("objDanMisDetails", []):
        row = {}

        row["a1"] = f"{item.get('slSanctDt', '')}/{item.get('totalSancAmt', '')}"
        row["a2"] = item.get("schemeName", "")
        row["a3"] = item.get("slAcctNo", "")
        row["a4"] = f"{item.get('disbDt', '')}/{item.get('terminalRepDate', '')}"
        row["a5"] = f"{item.get('totalAmtDisb', '')}/{item.get('osAmount', '')}"
        row["a6"] = item.get("roi", "")
        row["a7"] = f"{item.get('resetClause', '')}/{item.get('nextResetDate', '')}"
        row["a8"] = item.get("conduct", "")

        sanction_rows.append(row)

    mis_json["sanction"] = sanction_rows

    mjso = dict(mis_json)

    # print(mjso)
    # ------------------------------------------------------------
    # Recursive helper to collect all "result" values from LLM JSON
    # ------------------------------------------------------------

    def collect_results(obj, output):
        """
        Recursively walks through the LLM JSON.

        If it finds:
            {
                "result": {
                    "Some Key": {
                        "value": "...",
                        "remarks": "...",
                        "notes": "..."
                    }
                }
            }

        It converts "Some Key" -> "some_key" and stores the value.
        """

        if isinstance(obj, dict):
            if "result" in obj and isinstance(obj["result"], dict):
                for key, value in obj["result"].items():
                    new_key = key.replace(" ", "_").lower()
                    new_key = new_key.replace("{","").replace("}","")

                    if isinstance(value, dict):
                        if "value" in value:
                            #new change
                            output[new_key] = value["value"] if (value["value"] or value["value"] == 0) else "NOT_FOUND"
                        elif "remarks" in value:
                            output[new_key] = value["remarks"] if value["remarks"] else "NOT_FOUND"
                        elif "notes" in value:
                            output[new_key] = value["notes"] if value["notes"] else "NOT_FOUND"
                        else:
                            if isinstance(value,dict):
                                for i,j in value.items():
                                    #new change
                                    output[f"{new_key}_{i}"] = j if (j or j == 0) else "NOT_FOUND"
                            else:
                                output[new_key] = "NOT_FOUND"
                    else:
                        output[new_key] = value if value else "NOT_FOUND"

            for value in obj.values():
                collect_results(value, output)

        elif isinstance(obj, list):
            for item in obj:
                collect_results(item, output)

    # ------------------------------------------------------------
    # Collect LLM scalar outputs
    # ------------------------------------------------------------

    merged = defaultdict(list)
    collect_results(data, merged)

    ljso = dict(merged)
    
    with open('testx.json','w') as f:
        json.dump(ljso,f,indent=4)
    # ------------------------------------------------------------
    # Merge LLM and MIS values
    # Python 3.9+ dictionary union syntax
    # ------------------------------------------------------------

    result = ljso | mjso

    # ------------------------------------------------------------
    # Separate scalar values and table values
    # ------------------------------------------------------------

    ntbl = {}
    tbl = {}

    for key, value in result.items():
        if isinstance(value, list):
            tbl[key] = value
        else:
            ntbl[key] = value

    # ------------------------------------------------------------
    # Replace scalar placeholders in HTML template
    # Example:
    #   {{ borrower_name }} -> actual value
    # ------------------------------------------------------------

    for k in list(tbl.keys()):
        if not tbl[k]:
            ntbl[k] = ""
            del tbl[k]
    #new change
    for k, v in ntbl.items():
        if (not v and v != 0) or (isinstance(v, str) and v.strip() == ""):
            ntbl[k] = "NOT_FOUND"

            
    with open("result.json","w",encoding='utf-8') as f:
        json.dump(result,f,indent=4)
    
    
    result_html = re.sub(
        r"\{\{\s*(.*?)\s*\}\}",
        lambda match: (
            str(ntbl[key]) if (key := match.group(1).strip().lower()) in ntbl
            else match.group(0)
        ),
        html_template,
    )

    # ------------------------------------------------------------
    # Mapping between JSON table keys and HTML table IDs
    # ------------------------------------------------------------

    table_id_mapping = {
        "product_concentration": "product",
        "geographic_concentration": "geo",
        "portfolio_mix": "portfolio",
        "borrowing_mix": "borrow",
        "top_5_lenders": "lend",
        "secured_vs_unsecured_mix": "secure",
        "sanction": "past-live",
    }

    # ------------------------------------------------------------
    # Convert table dictionaries into row lists
    # ------------------------------------------------------------

    table_rows_by_html_id = defaultdict(list)

    for json_key, html_table_id in table_id_mapping.items():
        
        for row_obj in tbl.get(json_key, []):
            if isinstance(row_obj, dict):
                table_rows_by_html_id[html_table_id].append(list(row_obj.values()))
            else:
                table_rows_by_html_id[html_table_id].append(row_obj)

    final_table_data = dict(table_rows_by_html_id)

    # print(final_table_data)
    # print(final_table_data)
    # ------------------------------------------------------------
    # Parse HTML and append dynamic rows into matching tables
    # ------------------------------------------------------------

    soup = BeautifulSoup(result_html, "html.parser")

    for html_table_id, rows in final_table_data.items():
        table = soup.find("table", id=html_table_id)

        if not table:
            # Table ID not found in HTML template.
            # Intentionally skipping instead of failing.
            continue

        # Prefer tbody if present, otherwise append directly to table.
        append_target = table.find("tbody") or table
        
        for row_values in rows:
            tr = soup.new_tag("tr")

            for cell_value in row_values:
                td = soup.new_tag("td")
                td.string = str(cell_value)
                tr.append(td)

            append_target.append(tr)
            
            
    result_html = str(soup)

    # ------------------------------------------------------------
    # Replace any still-unresolved placeholders with rm_input
    # These will be highlighted later in DOCX.
    # ------------------------------------------------------------

    result_html = re.sub(
        r"\{\{\s*(.*?)\s*\}\}",
        r"<i>&lt;RM input&gt;</i>",
        result_html,
    )

    return result_html.strip().replace('NOT_FOUND','RM to check and update').replace('Not_found','RM to check and update').replace('not_found','RM to check and update').replace('Not Available','RM to check and update').replace('Not Found', 'RM to check and update')


# ============================================================
# BASIC HTML CLEANING
# ============================================================

def clean_html_before_parse(html: str) -> str:
    """
    Pre-cleans HTML before BeautifulSoup parsing.

    Important:
        - Does not alter template variables like {{ variable_name }}.
    """

    html = html.replace("```", "").replace("``", "")

    html = re.sub(r"<!--.*?-->", "", html, flags=re.DOTALL)

    html = re.sub(
        r'\s(width|height)\s*=\s*["\'][^"\']*["\']',
        "",
        html,
        flags=re.IGNORECASE,
    )

    html = re.sub(
        r'\s(width|height)\s*=\s*[^\s>]+',
        "",
        html,
        flags=re.IGNORECASE,
    )

    return html.replace('NOT_FOUND','RM to check and update').replace('Not_found','RM to check and update').replace('not_found','RM to check and update').replace('Not Available','RM to check and update')


def parse_inline_style(tag: Tag) -> dict:
    """
    Parses inline CSS style into a dictionary.
    Example:
        style="font-weight:bold; color:#000000"
    becomes:
        {"font-weight": "bold", "color": "#000000"}
    """

    if not isinstance(tag, Tag):
        return {}

    style = tag.get("style", "")
    result = {}

    for item in style.split(";"):
        if ":" not in item:
            continue

        key, value = item.split(":", 1)
        key = key.strip().lower()
        value = value.strip()

        if key and value:
            result[key] = value

    return result


def strip_layout_styles(soup: BeautifulSoup) -> None:
    """
    Removes Word-hostile layout styles from HTML.

    Removes:
        - width
        - height
        - margins
        - padding
        - table-layout
        - border-collapse
        - border styles

    Keeps useful styles:
        - text-align
        - font-size
        - font-weight
        - color
        - background-color
    """

    remove_props = {
        "width",
        "height",
        "min-width",
        "max-width",
        "min-height",
        "max-height",
        "table-layout",
        "border-collapse",
        "border-spacing",
        "padding",
        "padding-top",
        "padding-right",
        "padding-bottom",
        "padding-left",
        "margin-left",
        "margin-right",
        "margin-top",
        "margin-bottom",
        "text-indent",
        "border",
        "border-top",
        "border-right",
        "border-bottom",
        "border-left",
    }

    # Remove comments.
    for comment in soup.find_all(string=lambda text: isinstance(text, Comment)):
        comment.extract()

    # Clean all tags.
    for tag in soup.find_all(True):
        # Remove direct width/height attributes.
        for attr in ["width", "height"]:
            if attr in tag.attrs:
                del tag.attrs[attr]

        # Remove malformed attributes occasionally generated by HTML templates.
        for attr in list(tag.attrs.keys()):
            if attr.lower() in {"border-collapse", "collapse;"}:
                del tag.attrs[attr]

        style = tag.get("style")

        if not style:
            continue

        new_styles = []

        for item in style.split(";"):
            if ":" not in item:
                continue

            key, value = item.split(":", 1)
            key = key.strip().lower()
            value = value.strip()

            if not key or not value:
                continue

            if key in remove_props:
                continue

            new_styles.append(f"{key}: {value}")

        if new_styles:
            tag["style"] = "; ".join(new_styles)
        else:
            del tag.attrs["style"]

    # Remove empty spans and white tracker spans.
    for span in list(soup.find_all("span")):
        text = span.get_text().replace("\xa0", "").strip()
        style = span.get("style", "").lower()

        is_white_tracker = (
            "color: #ffffff" in style
            or "color:#ffffff" in style
            or "color: white" in style
        )

        if not text or is_white_tracker:
            span.unwrap()


# ============================================================
# DOCX XML HELPERS
# ============================================================

def set_cell_border(
    cell,
    color: str = OUTPUT_TABLE_BORDER_COLOR,
    size: str = OUTPUT_TABLE_BORDER_SIZE,
) -> None:
    """
    Applies border to a DOCX table cell.
    """

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    borders = tc_pr.first_child_found_in("w:tcBorders")

    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ["top", "left", "bottom", "right"]:
        edge_tag = f"w:{edge}"
        element = borders.find(qn(edge_tag))

        if element is None:
            element = OxmlElement(edge_tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str = HEADER_FILL_COLOR) -> None:
    """
    Applies background shading to a DOCX table cell.
    """

    tc_pr = cell._tc.get_or_add_tcPr()

    shading = tc_pr.first_child_found_in("w:shd")

    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)

    shading.set(qn("w:fill"), fill)


def clear_cell(cell) -> None:
    """
    Removes all default paragraphs from a DOCX cell.
    """

    for paragraph in list(cell.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)


def apply_table_formatting(table) -> None:
    """
    Applies common formatting to DOCX tables.
    """

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_border(cell)


def set_document_defaults(doc: Document) -> None:
    """
    Sets global document margins and default font.
    """

    section = doc.sections[0]

    section.top_margin = Inches(PAGE_MARGIN_TOP)
    section.bottom_margin = Inches(PAGE_MARGIN_BOTTOM)
    section.left_margin = Inches(PAGE_MARGIN_LEFT)
    section.right_margin = Inches(PAGE_MARGIN_RIGHT)

    normal = doc.styles["Normal"]
    normal.font.name = DEFAULT_FONT
    normal.font.size = Pt(DEFAULT_FONT_SIZE_PT)


# ============================================================
# STYLE EXTRACTION
# ============================================================

def html_color_to_rgb(value: str):
    """
    Converts '#RRGGBB' HTML color to python-docx RGBColor.
    """

    if not value:
        return None

    value = value.strip()

    match = re.match(r"#([0-9a-fA-F]{6})", value)

    if not match:
        return None

    hex_value = match.group(1)

    return RGBColor(
        int(hex_value[0:2], 16),
        int(hex_value[2:4], 16),
        int(hex_value[4:6], 16),
    )


def get_font_size_pt(tag: Tag):
    """
    Extracts font-size from inline style and converts it to points.
    Supports:
        - pt
        - px
    """

    styles = parse_inline_style(tag)
    value = styles.get("font-size")

    if not value:
        return None

    value = value.lower().strip()

    try:
        if value.endswith("pt"):
            return float(value[:-2].strip())

        if value.endswith("px"):
            return float(value[:-2].strip()) * 0.75

    except Exception:
        return None

    return None


def get_alignment(tag: Tag):
    """
    Converts HTML text-align into python-docx paragraph alignment.
    """

    styles = parse_inline_style(tag)
    align = styles.get("text-align", "").lower()

    if align == "center":
        return WD_ALIGN_PARAGRAPH.CENTER

    if align == "right":
        return WD_ALIGN_PARAGRAPH.RIGHT

    if align == "justify":
        return WD_ALIGN_PARAGRAPH.JUSTIFY

    if align == "left":
        return WD_ALIGN_PARAGRAPH.LEFT

    return None


def tag_is_bold(tag: Tag) -> bool:
    """
    Detects whether a tag should render as bold.
    """

    if not isinstance(tag, Tag):
        return False

    if tag.name in {"b", "strong"}:
        return True

    styles = parse_inline_style(tag)
    font_weight = styles.get("font-weight", "").lower()

    return font_weight in {"bold", "700", "800", "900"}


def tag_is_italic(tag: Tag) -> bool:
    """
    Detects whether a tag should render as italic.
    """

    if not isinstance(tag, Tag):
        return False

    if tag.name in {"i", "em"}:
        return True

    styles = parse_inline_style(tag)

    return styles.get("font-style", "").lower() == "italic"


def tag_is_underline(tag: Tag) -> bool:
    """
    Detects whether a tag should render as underlined.
    """

    if not isinstance(tag, Tag):
        return False

    if tag.name == "u":
        return True

    styles = parse_inline_style(tag)

    return "underline" in styles.get("text-decoration", "").lower()


def get_background_color(tag: Tag):
    """
    Extracts background color from a tag.
    Returns hex color without '#'.
    """

    styles = parse_inline_style(tag)
    bg = styles.get("background-color") or styles.get("background")

    if not bg:
        return None

    match = re.match(r"#([0-9a-fA-F]{6})", bg.strip())

    if match:
        return match.group(1).upper()

    return None


# ============================================================
# TEXT RENDERING
# ============================================================

def add_run_with_context(paragraph, text: str, context: list[Tag]) -> None:
    """
    Adds a run to a paragraph using inherited formatting from context tags.
    """

    if text is None:
        return

    text = str(text).replace("\xa0", " ")

    if text == "":
        return

    run = paragraph.add_run(text)

    run.font.name = DEFAULT_FONT
    run.font.size = Pt(DEFAULT_FONT_SIZE_PT)

    for tag in context:
        if tag_is_bold(tag):
            run.bold = True

        if tag_is_italic(tag):
            run.italic = True

        if tag_is_underline(tag):
            run.underline = True

        font_size = get_font_size_pt(tag)

        if font_size:
            run.font.size = Pt(font_size)

        styles = parse_inline_style(tag)
        color = html_color_to_rgb(styles.get("color", ""))

        if color:
            run.font.color.rgb = color


def render_inline(paragraph, node, context=None) -> None:
    """
    Recursively renders inline HTML nodes into a DOCX paragraph.
    """

    if context is None:
        context = []

    if isinstance(node, NavigableString):
        add_run_with_context(paragraph, str(node), context)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run().add_break()
        return

    new_context = context + [node]

    for child in node.children:
        render_inline(paragraph, child, new_context)


def add_paragraph_from_html(container, p_tag: Tag):
    """
    Adds a DOCX paragraph from an HTML <p> tag.
    """

    paragraph = container.add_paragraph()

    alignment = get_alignment(p_tag)

    if alignment is not None:
        paragraph.alignment = alignment

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)

    for child in p_tag.children:
        render_inline(paragraph, child, [p_tag])

    return paragraph


def add_text_paragraph(container, text: str):
    """
    Adds a plain text paragraph.
    """

    paragraph = container.add_paragraph()

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)

    add_run_with_context(paragraph, text, [])

    return paragraph


# ============================================================
# TABLE EXTRACTION
# ============================================================

def get_direct_rows(table_tag: Tag) -> list[Tag]:
    """
    Gets rows belonging directly to the given table.

    Handles:
        <table>
            <tbody>
                <tr>...</tr>
            </tbody>
            <tr>...</tr>
        </table>
    """

    rows = []

    for child in table_tag.children:
        if not isinstance(child, Tag):
            continue

        if child.name == "tr":
            rows.append(child)

        elif child.name in {"tbody", "thead", "tfoot"}:
            for row in child.children:
                if isinstance(row, Tag) and row.name == "tr":
                    rows.append(row)

    return rows


def get_direct_cells(row_tag: Tag) -> list[Tag]:
    """
    Gets direct td/th cells from a row.
    """

    return [
        child
        for child in row_tag.children
        if isinstance(child, Tag) and child.name in {"td", "th"}
    ]


def safe_int(value, default=1) -> int:
    """
    Safely converts a value to int.
    Ensures minimum value is 1.
    """

    try:
        return max(1, int(str(value).strip()))
    except Exception:
        return default


def build_table_grid(table_tag: Tag):
    """
    Builds a logical table grid considering colspan and rowspan.

    Returns:
        grid, row_count, col_count
    """

    rows = get_direct_rows(table_tag)

    grid = []
    occupied = {}

    max_cols = 0

    for r_idx, row in enumerate(rows):
        cells = get_direct_cells(row)

        if not cells:
            grid.append([])
            continue

        if len(grid) <= r_idx:
            grid.append([])

        c_idx = 0

        for cell in cells:
            while occupied.get((r_idx, c_idx)):
                c_idx += 1

            rowspan = safe_int(cell.get("rowspan", 1))
            colspan = safe_int(cell.get("colspan", 1))

            while len(grid[r_idx]) <= c_idx:
                grid[r_idx].append(None)

            entry = {
                "tag": cell,
                "origin": True,
                "rowspan": rowspan,
                "colspan": colspan,
            }

            grid[r_idx][c_idx] = entry

            for rr in range(r_idx, r_idx + rowspan):
                while len(grid) <= rr:
                    grid.append([])

                for cc in range(c_idx, c_idx + colspan):
                    occupied[(rr, cc)] = True

                    while len(grid[rr]) <= cc:
                        grid[rr].append(None)

                    if not (rr == r_idx and cc == c_idx):
                        grid[rr][cc] = {
                            "tag": cell,
                            "origin": False,
                            "master": (r_idx, c_idx),
                        }

            c_idx += colspan
            max_cols = max(max_cols, c_idx)

    for row in grid:
        max_cols = max(max_cols, len(row))

    if max_cols == 0:
        max_cols = 1

    for row in grid:
        while len(row) < max_cols:
            row.append(None)

    return grid, len(grid), max_cols


# ============================================================
# TABLE RENDERING
# ============================================================

def render_table(container, table_tag: Tag):
    """
    Renders an HTML table into DOCX.
    Supports:
        - rowspan
        - colspan
        - nested tables
        - th header shading
        - background color
    """

    grid, row_count, col_count = build_table_grid(table_tag)

    if row_count == 0 or col_count == 0:
        return None

    # Remove fully empty rows caused by malformed HTML.
    meaningful_rows = []

    for row in grid:
        has_origin = any(cell and cell.get("origin") for cell in row)

        if has_origin:
            meaningful_rows.append(row)

    if not meaningful_rows:
        return None

    grid = meaningful_rows
    row_count = len(grid)

    docx_table = container.add_table(rows=row_count, cols=col_count)
    apply_table_formatting(docx_table)

    for r_idx, row in enumerate(grid):
        for c_idx, entry in enumerate(row):
            if not entry or not entry.get("origin"):
                continue

            html_cell = entry["tag"]
            rowspan = entry.get("rowspan", 1)
            colspan = entry.get("colspan", 1)

            start_cell = docx_table.cell(r_idx, c_idx)

            end_r = min(row_count - 1, r_idx + rowspan - 1)
            end_c = min(col_count - 1, c_idx + colspan - 1)

            target_cell = start_cell

            if end_r != r_idx or end_c != c_idx:
                try:
                    target_cell = start_cell.merge(docx_table.cell(end_r, end_c))
                except Exception:
                    target_cell = start_cell

            clear_cell(target_cell)

            if html_cell.name == "th":
                set_cell_shading(target_cell, HEADER_FILL_COLOR)

            bg = get_background_color(html_cell)

            if bg:
                set_cell_shading(target_cell, bg)

            render_cell_content(target_cell, html_cell)

    return docx_table


def render_cell_content(cell, html_cell: Tag) -> None:
    """
    Renders the content inside one HTML table cell into one DOCX cell.
    """

    added = False

    for child in html_cell.children:
        if isinstance(child, NavigableString):
            text = str(child).replace("\xa0", " ")

            if text.strip():
                paragraph = cell.add_paragraph()

                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)

                alignment = get_alignment(html_cell)

                if alignment is not None:
                    paragraph.alignment = alignment

                add_run_with_context(paragraph, text, [html_cell])
                added = True

        elif isinstance(child, Tag):
            if child.name == "p":
                add_paragraph_from_html(cell, child)
                added = True

            elif child.name == "table":
                render_table(cell, child)
                added = True

            elif child.name == "br":
                cell.add_paragraph()
                added = True

            elif child.name in {"b", "strong", "span", "i", "em", "u"}:
                paragraph = cell.add_paragraph()

                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)

                alignment = get_alignment(html_cell)

                if alignment is not None:
                    paragraph.alignment = alignment

                render_inline(paragraph, child, [html_cell])
                added = True

            else:
                # Render nested direct block children if present.
                direct_blocks = [
                    x
                    for x in child.children
                    if isinstance(x, Tag) and x.name in {"p", "table"}
                ]

                if direct_blocks:
                    for block in direct_blocks:
                        if block.name == "p":
                            add_paragraph_from_html(cell, block)
                            added = True

                        elif block.name == "table":
                            render_table(cell, block)
                            added = True

                else:
                    text = child.get_text(" ", strip=False)

                    if text.strip():
                        paragraph = cell.add_paragraph()

                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(3)

                        alignment = get_alignment(child) or get_alignment(html_cell)

                        if alignment is not None:
                            paragraph.alignment = alignment

                        render_inline(paragraph, child, [html_cell, child])
                        added = True

    if not added:
        cell.add_paragraph("")


# ============================================================
# BODY RENDERING
# ============================================================

def render_body(doc: Document, body: Tag) -> None:
    """
    Renders parsed HTML body into a DOCX document.
    """

    for child in body.children:
        if isinstance(child, NavigableString):
            text = str(child).replace("\xa0", " ").strip()

            if text:
                add_text_paragraph(doc, text)

            continue

        if not isinstance(child, Tag):
            continue

        if child.name == "p":
            add_paragraph_from_html(doc, child)

        elif child.name == "table":
            render_table(doc, child)

        elif child.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = safe_int(child.name[1], 1)

            paragraph = doc.add_heading(level=min(level, 4))
            render_inline(paragraph, child, [child])

        elif child.name == "br":
            doc.add_paragraph("")

        else:
            # For unknown wrappers, render only direct p/table children.
            found_block = False

            for sub in child.children:
                if isinstance(sub, Tag) and sub.name == "p":
                    add_paragraph_from_html(doc, sub)
                    found_block = True

                elif isinstance(sub, Tag) and sub.name == "table":
                    render_table(doc, sub)
                    found_block = True

            if not found_block:
                text = child.get_text(" ", strip=True)

                if text:
                    add_text_paragraph(doc, text)


# ============================================================
# FINAL POST-PROCESSING: HIGHLIGHT rm_input
# ============================================================

def highlight_rm_input_in_docx(
    doc: Document,
    target: str,
    fill_rgb: str,
) -> None:
    """
    Highlights every exact occurrence of target text using custom RGB shading.

    Args:
        doc: python-docx Document object
        target: exact text to highlight
        fill_rgb: RGB hex color without '#', for example:
                  - 'FFFF00' yellow
                  - 'FFEB9C' light yellow
                  - 'FFD966' gold
                  - 'D9EAD3' light green
                  - 'D9EAF7' light blue

    Works in:
        - normal paragraphs
        - top-level tables
        - nested tables
    """

    fill_rgb = fill_rgb.replace("#", "").upper()

    def make_run_like(original_run_element, text: str, highlight: bool = False):
        """
        Creates a new XML run based on an existing run's formatting.
        Applies custom RGB background shading when highlight=True.
        """

        new_run = OxmlElement("w:r")

        original_rpr = original_run_element.find(qn("w:rPr"))

        if original_rpr is not None:
            new_rpr = deepcopy(original_rpr)
        else:
            new_rpr = OxmlElement("w:rPr")

        if highlight:
            # Remove existing Word highlight if present.
            for old_highlight in list(new_rpr.findall(qn("w:highlight"))):
                new_rpr.remove(old_highlight)

            # Remove existing shading if present.
            for old_shading in list(new_rpr.findall(qn("w:shd"))):
                new_rpr.remove(old_shading)

            # Add custom RGB shading.
            shading_element = OxmlElement("w:shd")
            shading_element.set(qn("w:val"), "clear")
            shading_element.set(qn("w:color"), "auto")
            shading_element.set(qn("w:fill"), fill_rgb)
            new_rpr.append(shading_element)

        new_run.append(new_rpr)

        text_element = OxmlElement("w:t")
        text_element.set(qn("xml:space"), "preserve")
        text_element.text = text

        new_run.append(text_element)

        return new_run

    def highlight_in_paragraph(paragraph):
        """
        Splits runs containing target text and highlights only target pieces.
        """

        for run in list(paragraph.runs):
            if target not in run.text:
                continue

            original_text = run.text
            original_run_element = run._r
            parent = original_run_element.getparent()
            insert_index = parent.index(original_run_element)

            parts = re.split(f"({re.escape(target)})", original_text)

            new_runs = []

            for part in parts:
                if part == "":
                    continue

                new_runs.append(
                    make_run_like(
                        original_run_element,
                        part,
                        highlight=(part == target),
                    )
                )

            for offset, new_run in enumerate(new_runs):
                parent.insert(insert_index + offset, new_run)

            parent.remove(original_run_element)

    def iter_all_paragraphs_from_table(table):
        """
        Yields all paragraphs inside a table, including nested tables.
        """

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

                for nested_table in cell.tables:
                    yield from iter_all_paragraphs_from_table(nested_table)

    # Highlight top-level paragraphs.
    for paragraph in doc.paragraphs:
        highlight_in_paragraph(paragraph)

    # Highlight table and nested-table paragraphs.
    for table in doc.tables:
        for paragraph in iter_all_paragraphs_from_table(table):
            highlight_in_paragraph(paragraph)

# Adds a red disclaimer footer at the end of each page.
def add_ai_disclaimer_footer(doc: Document) -> None:
    disclaimer_text = (
        "*Fields are AI-generated; to be verified and used by the Bank Officer."
    )

    for section in doc.sections:
        footer = section.footer

        # Use existing first footer paragraph if available, else create one
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
            paragraph.clear()
        else:
            paragraph = footer.add_paragraph()

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = paragraph.add_run(disclaimer_text)
        run.font.name = DEFAULT_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.italic = True

# ============================================================
# FUNCTION 2:
# HTML STRING -> DOCX BYTES
# ============================================================

def html_to_docx_bytes(input_html: str) -> bytes:
    """
    Takes:
        - input_html: complete HTML string

    Returns:
        - DOCX file as bytes

    Notes:
        - Does not read from disk.
        - Does not save to disk.
        - Suitable for FastAPI StreamingResponse / Response.
    """

    # ------------------------------------------------------------
    # Keep count of original template variables for optional checking.
    # This is useful if you later want diagnostics.
    # ------------------------------------------------------------

    original_vars = re.findall(r"\{\{[^}]+\}\}", input_html)
    original_var_counter = Counter(original_vars)

    # ------------------------------------------------------------
    # Clean and parse HTML.
    # ------------------------------------------------------------

    cleaned_html = clean_html_before_parse(input_html)

    soup = BeautifulSoup(cleaned_html, "lxml")

    strip_layout_styles(soup)

    body = soup.body if soup.body else soup

    # ------------------------------------------------------------
    # Create DOCX in memory.
    # ------------------------------------------------------------

    doc = Document()

    set_document_defaults(doc)

    add_ai_disclaimer_footer(doc)

    render_body(doc, body)

    # ------------------------------------------------------------
    # Highlight unresolved fields.
    # ------------------------------------------------------------

    highlight_rm_input_in_docx(doc, "<RM input>", " FFFACD")
    highlight_rm_input_in_docx(doc, "RM to check and update", "D9EAF7")
    
    # ------------------------------------------------------------
    # Optional visibility check.
    # Kept here as comments/logical block but not printed,
    # because this function should only return bytes.
    # ------------------------------------------------------------

    visible_text = []

    for paragraph in doc.paragraphs:
        visible_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    visible_text.append(paragraph.text)

    output_text = "\n".join(visible_text)
    output_vars = re.findall(r"\{\{[^}]+\}\}", output_text)
    output_var_counter = Counter(output_vars)

    missing = list((original_var_counter - output_var_counter).elements())

    # You may log `missing` outside this function if needed.
    # Here we intentionally do nothing to keep function side-effect free.
    _ = missing

    # ------------------------------------------------------------
    # Save DOCX to BytesIO instead of filesystem.
    # ------------------------------------------------------------

    output_stream = BytesIO()

    doc.save(output_stream)

    output_stream.seek(0)

    return output_stream.getvalue()


# ============================================================
# FUNCTION 3:
# JSON INPUT + HTML TEMPLATE -> DOCX BYTES
# ============================================================

def json_to_docx_bytes(html_template: str, json_data: dict) -> bytes:
    """
    Combines Function 1 and Function 2.

    Takes:
        - html_template: raw HTML template string
        - json_data: input JSON dictionary

    Returns:
        - final DOCX file as bytes

    Flow:
        JSON + HTML template
            -> filled HTML string
            -> DOCX bytes
    """

    filled_html = json_to_filled_html_string(
        html_template=html_template,
        json_data=json_data,
    )

    docx_bytes = html_to_docx_bytes(
        input_html=filled_html.replace('NOT_FOUND','RM to check and update').replace('Not_found','RM to check and update').replace('not_found','RM to check and update').replace('Not Available','RM to check and update'),
    )

    return docx_bytes

def json_to_docx_bytes_from_default_template(json_data: dict) -> bytes:
    """
    Uses backend-stored input.html as the fixed HTML template.

    User only needs to provide:
        - json_data

    Returns:
        - DOCX file as bytes

    Required backend file:
        - input.html must exist in the same directory where this Python file runs.
    """

    # Read fixed HTML template from backend
    html_template = Path(INPUT_HTML_FILE).read_text(
        encoding="utf-8",
        errors="replace"
    )

    # Reuse existing combined function
    return json_to_docx_bytes(
        html_template=html_template,
        json_data=json_data,
    )
    
    
with open("axis.json", "r", encoding="utf-8") as f:
    json_data = json.load(f)

docx_bytes = json_to_docx_bytes_from_default_template(
    json_data=json_data,
 )

with open("converted_output.docx", "wb") as f:
    f.write(docx_bytes)
